from docx import Document

# Данные
results = [
    ("classic", 1.3717),
    ("global_shuffle", 0.7088),
    ("by_distance_shuffle", 1.4930),
    ("stride_window", 1.3928),
    ("random_chunks", 0.3460),
]

# Создание документа Word
doc = Document()
doc.add_heading("📋 Сводка по всем вариантам подачи данных", level=1)

# Таблица
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Заголовки
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Вариант подачи данных'
hdr_cells[1].text = 'MAE (м)'

# Добавление данных
for method, mae in results:
    row_cells = table.add_row().cells
    row_cells[0].text = method
    row_cells[1].text = f"{mae:.4f}"

# Сохранение
doc_path = "newFile.docx"
doc.save(doc_path)

doc_path